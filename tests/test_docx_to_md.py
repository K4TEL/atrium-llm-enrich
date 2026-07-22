"""
tests/test_docx_to_md.py
========================
Tests for api_util/docx_to_md.py — DOCX → visually-rich Markdown. A tiny .docx
is built in-process with python-docx (hermetic, no committed binaries), so the
whole suite is skipped cleanly wherever python-docx isn't installed.
"""

import pytest

from api_util import docx_to_md
from api_util.docx_to_md import _heading_prefix

pytestmark = pytest.mark.skipif(not docx_to_md.docx_available(), reason="python-docx not installed")


@pytest.fixture
def sample_docx(tmp_path):
    """A DOCX exercising headings, emphasis, font/colour/highlight, alignment,
    a table, a header/footer and a hard page break."""
    import docx
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_COLOR_INDEX
    from docx.shared import Pt, RGBColor

    d = docx.Document()
    d.sections[0].header.paragraphs[0].text = "Report header"
    d.sections[0].footer.paragraphs[0].text = "Report footer"

    d.add_heading("Hlavni nadpis", level=1)

    p = d.add_paragraph()
    p.add_run("tucne").bold = True
    p.add_run(" ")
    p.add_run("kurziva").italic = True
    p.add_run(" ")
    p.add_run("skrtnute").font.strike = True

    styled = d.add_paragraph()
    r = styled.add_run("barevne")
    r.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    r.font.size = Pt(14)
    r.font.name = "Courier New"
    r2 = styled.add_run(" zvyraznene")
    r2.font.highlight_color = WD_COLOR_INDEX.YELLOW

    centered = d.add_heading("Centrovany", level=2)
    centered.alignment = WD_ALIGN_PARAGRAPH.CENTER

    t = d.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "A"
    t.cell(0, 1).text = "B|pipe"
    t.cell(1, 0).text = "1"
    t.cell(1, 1).text = "2"

    d.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    d.add_paragraph("Text na druhe strance.")

    dest = tmp_path / "CTX999.docx"
    d.save(str(dest))
    return dest


def test_heading_prefix_mapping():
    assert _heading_prefix("Heading 1") == "#"
    assert _heading_prefix("Heading 3") == "###"
    assert _heading_prefix("Title") == "#"
    assert _heading_prefix("Normal") is None
    assert _heading_prefix(None) is None


def test_convert_title_from_doc_id(sample_docx):
    md = docx_to_md.convert(sample_docx)
    assert md.startswith("# CTX999")


def test_convert_page_sections_and_break(sample_docx):
    md = docx_to_md.convert(sample_docx)
    assert "## Page 1" in md
    assert "<!-- PAGE_BREAK: pg_2 -->" in md
    assert "## Page 2" in md
    assert md.index("## Page 2") > md.index("## Page 1")
    assert "Text na druhe strance." in md
    # the second page's body comes after the page-2 heading
    assert md.index("Text na druhe strance.") > md.index("## Page 2")


def test_convert_emphasis_forms(sample_docx):
    md = docx_to_md.convert(sample_docx)
    assert "**tucne**" in md
    assert "*kurziva*" in md
    assert "~~skrtnute~~" in md


def test_convert_font_and_style_cues(sample_docx):
    md = docx_to_md.convert(sample_docx)
    assert '<!-- FONT: size=14pt, family="Courier New" -->' in md
    assert "<!-- STYLE: color=#FF0000 -->" in md
    assert "<!-- STYLE: highlight=yellow -->" in md


def test_convert_alignment_wraps_heading(sample_docx):
    md = docx_to_md.convert(sample_docx)
    assert '<div align="center">' in md
    assert "## Centrovany" in md


def test_convert_table_as_gfm(sample_docx):
    md = docx_to_md.convert(sample_docx)
    assert "| A | B\\|pipe |" in md
    assert "| --- | --- |" in md


def test_convert_headers_and_footers(sample_docx):
    md = docx_to_md.convert(sample_docx)
    assert "<!-- HEADER_START -->" in md
    assert "Report header" in md
    assert "<!-- FOOTER_START -->" in md
    assert "Report footer" in md


def test_convert_geometry_cues(sample_docx):
    md = docx_to_md.convert(sample_docx)
    assert "<!-- DOC_META:" in md
    assert "<!-- LAYOUT_MARGIN:" in md


def test_convert_missing_lib_raises(monkeypatch, sample_docx):
    monkeypatch.setattr(docx_to_md, "docx_available", lambda: False)
    with pytest.raises(docx_to_md.DocxNotInstalled):
        docx_to_md.convert(sample_docx)
