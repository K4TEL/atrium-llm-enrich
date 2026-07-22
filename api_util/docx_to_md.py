"""
api_util/docx_to_md.py — DOCX → visually-rich Markdown.

Converts a Word ``.docx`` into the page-sectioned Markdown consumed by
``run_document_level()`` (llm_client_shared.py), enriched with the visual-layout
cues catalogued in ``layout_md.py`` (issue #10). DOCX is the *easy* input: it
carries a real Unicode text layer plus native styling metadata (fonts, colours,
highlights, alignment, styles → headings, tables, headers/footers, page-size /
margins), so most of the taxonomy can be recovered directly rather than inferred.

Uses **python-docx** (MIT), imported lazily so the base install never requires
it — mirroring ``flexiconv_convert.py``'s ``*_available()`` pattern. Output
matches ``xml_to_md.py``'s house style (``# <doc-id>`` title, ``## Page N``
sections) so the whole-document system prompt's page-locator instructions
resolve against real anchors.
"""

from __future__ import annotations

import argparse
import sys
from pathlib import Path
from typing import List, Optional

_repo_root = str(Path(__file__).resolve().parent.parent)
if _repo_root not in sys.path:
    sys.path.insert(0, _repo_root)

from api_util import layout_md as L  # noqa: E402
from api_util.teitok_read import doc_id_from_path  # noqa: E402

INSTALL_HINT = "python-docx is not installed. Please run: pip install -r requirements_docmd.txt"


class DocxNotInstalled(RuntimeError):
    pass


def docx_available() -> bool:
    """Whether python-docx can be imported, without raising."""
    try:
        import docx  # noqa: F401

        return True
    except ImportError:
        return False


# --------------------------------------------------------------------------- #
# Body walking — python-docx has no built-in in-order paragraph+table iterator.
# --------------------------------------------------------------------------- #
def _iter_block_items(parent):
    """Yield ``Paragraph`` and ``Table`` objects in document order.

    The standard python-docx recipe: walk the raw body/cell element and wrap
    each ``<w:p>`` / ``<w:tbl>`` child in its high-level object.
    """
    from docx.document import Document as _Document
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table, _Cell
    from docx.text.paragraph import Paragraph

    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        parent_elm = parent

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


# --------------------------------------------------------------------------- #
# Headings & runs
# --------------------------------------------------------------------------- #
def _heading_prefix(style_name: Optional[str]) -> Optional[str]:
    """Map a paragraph style name to a Markdown heading prefix, or None."""
    if not style_name:
        return None
    name = style_name.strip().lower()
    if name == "title":
        return "#"
    if name.startswith("heading"):
        digits = "".join(ch for ch in name if ch.isdigit())
        level = int(digits) if digits else 1
        return "#" * max(1, min(level, 6))
    if name == "subtitle":
        return "##"
    return None


def _color_hex(run) -> Optional[str]:
    """Explicit RGB font colour as ``#RRGGBB``, or None if inherited/auto."""
    try:
        rgb = run.font.color.rgb  # raises if color type isn't RGB
    except (AttributeError, KeyError, TypeError):
        return None
    return f"#{rgb}" if rgb is not None else None


def _highlight_name(run) -> Optional[str]:
    """Highlight colour as a lower-case name (``yellow``), or None."""
    try:
        hl = run.font.highlight_color
    except (AttributeError, KeyError):
        return None
    if hl is None:
        return None
    name = getattr(hl, "name", None)
    return name.lower() if name else None


def _run_font_cue(run) -> str:
    """FONT cue for a run's *explicit* size/family (inherited values omitted)."""
    size_pt = None
    try:
        if run.font.size is not None:
            size_pt = round(run.font.size.pt, 1)
    except (AttributeError, ValueError):
        size_pt = None
    family = run.font.name or None
    return L.font(size=size_pt, family=family)


def _run_to_md(run) -> str:
    """Render one run: text with inline emphasis + preceding FONT/STYLE cues.

    Cues are emitted only for *explicitly* set attributes — python-docx returns
    None for inherited properties — so unstyled body runs stay clean.
    """
    raw = run.text or ""
    if not raw.strip():
        return raw  # whitespace-only: keep spacing, add no markup

    # Hoist surrounding whitespace outside the emphasis markers so the Markdown
    # stays valid ("**Sonda** " not "**Sonda **", which CommonMark won't parse).
    core = raw.strip()
    lead = raw[: len(raw) - len(raw.lstrip())]
    trail = raw[len(raw.rstrip()) :]

    # Inline emphasis, layered inside-out.
    strike = bool(run.font.strike) if run.font.strike is not None else False
    if strike:
        core = L.strike(core)
    if run.underline:
        core = L.underline(core)
    if run.bold:
        core = L.bold(core)
    if run.italic:
        core = L.italic(core)

    cues = _run_font_cue(run) + L.style(color=_color_hex(run), highlight=_highlight_name(run))
    return f"{cues}{lead}{core}{trail}" if cues else f"{lead}{core}{trail}"


def _run_has_page_break(run) -> bool:
    """True if a run carries a hard ``<w:br w:type="page"/>``."""
    from docx.oxml.ns import qn

    for br in run._element.findall(qn("w:br")):
        if br.get(qn("w:type")) == "page":
            return True
    return False


def _run_image_md(run) -> Optional[str]:
    """Best-effort image placeholder for an inline drawing in a run."""
    from docx.oxml.ns import qn

    drawings = run._element.findall(".//" + qn("w:drawing"))
    if not drawings:
        return None
    alt, src = "image", "media/image"
    try:
        wp = "{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}"
        a = "{http://schemas.openxmlformats.org/drawingml/2006/main}"
        r = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}"
        doc_pr = run._element.find(".//" + wp + "docPr")
        if doc_pr is not None and doc_pr.get("name"):
            alt = doc_pr.get("name")
        blip = run._element.find(".//" + a + "blip")
        if blip is not None and blip.get(r + "embed"):
            rid = blip.get(r + "embed")
            part = run.part.related_parts.get(rid)
            if part is not None:
                src = str(part.partname).lstrip("/")
    except Exception:
        pass
    return L.image(alt, src)


def _collect_footnotes(document) -> dict:
    """Map footnote id -> text from the document's footnotes part (best-effort)."""
    notes: dict = {}
    try:
        from lxml import etree

        w = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        for rel in document.part.rels.values():
            if not str(rel.target_ref).endswith("footnotes.xml"):
                continue
            root = etree.fromstring(rel.target_part.blob)
            for fn in root.findall(w + "footnote"):
                fid = fn.get(w + "id")
                if fid is None or fid in {"-1", "0"}:  # separator / continuation
                    continue
                text = "".join(t.text or "" for t in fn.iter(w + "t")).strip()
                if text:
                    notes[fid] = text
            break
    except Exception:
        return {}
    return notes


def _paragraph_footnote_ids(paragraph) -> List[str]:
    """Footnote reference ids appearing in a paragraph, in order."""
    from docx.oxml.ns import qn

    ids = []
    for ref in paragraph._element.findall(".//" + qn("w:footnoteReference")):
        fid = ref.get(qn("w:id"))
        if fid is not None:
            ids.append(fid)
    return ids


def _paragraph_to_md(paragraph, referenced: set) -> str:
    """Render a paragraph: heading/alignment + inline runs, images, footnotes."""
    body = "".join(_run_image_md(run) or _run_to_md(run) for run in paragraph.runs).strip()

    # Append footnote markers for any references in this paragraph.
    for fid in _paragraph_footnote_ids(paragraph):
        body += L.footnote_ref(fid)
        referenced.add(fid)

    if not body:
        return ""

    prefix = _heading_prefix(getattr(paragraph.style, "name", None))
    if prefix:
        body = f"{prefix} {body}"

    align = None
    try:
        if paragraph.alignment is not None:
            align = paragraph.alignment.name.lower()  # LEFT/CENTER/RIGHT/JUSTIFY
    except AttributeError:
        align = None
    if align in {"center", "right", "justify"}:
        body = L.align_div(body, "center" if align == "justify" else align)

    return body


def _table_to_md(table) -> str:
    """Render a python-docx table as a GFM pipe table."""
    rows = []
    for row in table.rows:
        rows.append([cell.text for cell in row.cells])
    return L.md_table(rows, header=True)


# --------------------------------------------------------------------------- #
# Headers / footers / page geometry
# --------------------------------------------------------------------------- #
def _region_md(paragraphs, opener: str, closer: str) -> str:
    """Wrap non-empty header/footer paragraph text in a region cue block."""
    text = "\n".join(p.text.strip() for p in paragraphs if p.text and p.text.strip())
    if not text:
        return ""
    return f"{opener}\n{text}\n{closer}"


def _geometry_cues(section) -> List[str]:
    """DOC_META (size + orientation) and LAYOUT_MARGIN for a section."""
    cues: List[str] = []
    try:
        w_in = round(section.page_width.inches, 2)
        h_in = round(section.page_height.inches, 2)
        orient = getattr(section.orientation, "name", "portrait").lower()
        cues.append(L.doc_meta(size=f"{w_in}x{h_in}in", orientation=orient))
    except (AttributeError, TypeError):
        pass
    try:
        cues.append(
            L.layout_margin(
                top=f"{round(section.top_margin.inches, 2)}in",
                bottom=f"{round(section.bottom_margin.inches, 2)}in",
                left=f"{round(section.left_margin.inches, 2)}in",
                right=f"{round(section.right_margin.inches, 2)}in",
            )
        )
    except (AttributeError, TypeError):
        pass
    return cues


# --------------------------------------------------------------------------- #
# Public API
# --------------------------------------------------------------------------- #
def convert(path: str | Path) -> str:
    """Convert a ``.docx`` file to visually-rich, page-sectioned Markdown."""
    if not docx_available():
        raise DocxNotInstalled(INSTALL_HINT) from None
    import docx

    path = Path(path)
    document = docx.Document(str(path))
    footnotes = _collect_footnotes(document)
    referenced: set = set()

    parts: List[str] = [f"# {doc_id_from_path(path)}"]

    # Page geometry + repeating regions, once at the top.
    if document.sections:
        parts.extend(_geometry_cues(document.sections[0]))

    seen_regions: set = set()
    for section in document.sections:
        for paragraphs, opener, closer in (
            (section.header.paragraphs, L.header_start(), L.header_end()),
            (section.footer.paragraphs, L.footer_start(), L.footer_end()),
        ):
            block = _region_md(paragraphs, opener, closer)
            if block and block not in seen_regions:
                seen_regions.add(block)
                parts.append(block)

    page = 1
    parts.append(f"\n## Page {page}\n")

    for block in _iter_block_items(document):
        from docx.text.paragraph import Paragraph

        if isinstance(block, Paragraph):
            # Hard page break *before* this paragraph.
            if block.paragraph_format.page_break_before:
                page += 1
                parts.append(L.page_break(page))
                parts.append(f"\n## Page {page}\n")

            rendered = _paragraph_to_md(block, referenced)
            if rendered:
                parts.append(rendered)

            # Hard page break carried inside a run ends the current page.
            if any(_run_has_page_break(run) for run in block.runs):
                page += 1
                parts.append(L.page_break(page))
                parts.append(f"\n## Page {page}\n")
        else:  # Table
            table_md = _table_to_md(block)
            if table_md:
                parts.append(table_md)

    # Footnote definitions for markers actually emitted.
    defs = [L.footnote_def(fid, footnotes[fid]) for fid in referenced if fid in footnotes]
    if defs:
        parts.append("")
        parts.extend(sorted(defs))

    return "\n".join(parts).strip() + "\n"


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("input_file", type=Path)
    parser.add_argument(
        "--output", type=Path, default=None, help="Write to file instead of stdout."
    )
    args = parser.parse_args()

    if not args.input_file.exists():
        print(f"Input file not found: {args.input_file}", file=sys.stderr)
        sys.exit(1)

    try:
        rendered = convert(args.input_file)
    except DocxNotInstalled as exc:
        print(exc, file=sys.stderr)
        sys.exit(1)

    if args.output:
        args.output.write_text(rendered, encoding="utf-8")
        print(f"-> {args.output}")
    else:
        print(rendered)
